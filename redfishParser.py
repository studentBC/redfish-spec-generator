import requests
import sys
import ssl
import json
import docx
from docx import Document
from docx.shared import Inches
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.shared import RGBColor
ssl._create_default_https_context = ssl._create_unverified_context
visited=set()
# write document prepare
reload(sys)
sys.setdefaultencoding('utf-8')
document = Document()
mdoc = Document()
document.add_heading('Redfish Spec',0)
mdoc.add_heading('Missing url',0)
p = document.add_paragraph()
p.add_run('by benchin').bold = True
p.add_run()
dt={} #document table
dts={} #document table split uri into list
st={} #server BMC table
sts={} #server BMC table split uri into list


def write2Document (url, jobj, next):
    document.add_heading('GET ',level = 1)
    document.add_heading('Request ',level = 2)
    document.add_paragraph(url, style = 'ListBullet')
    table = document.add_table(rows = 1,cols = 4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Type'
    hdr_cells[2].text = 'Read Only'
    hdr_cells[3].text = 'Description'
    for key, value in jobj.items():
        #print(key,value)
        if key == '@odata.id' and value.find("/redfish/v1") > -1:
            #print("we add ",value)
            if value not in next:
                next.add(value)
            else:
                return False 
        elif key == 'Members':
            for members in value:
                if isinstance(value,list) and len(value):
                    if isinstance(value[0],dict) and value[0].has_key('@odata.id'):
                        next.add(value[0]['@odata.id'])
        elif key == 'ResourceBlocks':
            for members in value:
                if isinstance(value,list) and len(value):
                    if isinstance(value[0],dict) and value[0].has_key('@odata.id'):
                        next.add(value[0]['@odata.id'])
        elif type(value) is dict:
            for members in value.values():
                if type(members) is str and members.find("/redfish/v1") > -1:
                    next.add(members)

                        #print(value)
                #break
        if isinstance(value, dict):
            write2Document (url, value, next)
            continue

        new_cells = table.add_row().cells
        new_cells[0].text = str(key)
        if type(value) is list:
            new_cells[1].text = 'Array'
        elif type(value) is int:
            new_cells[1].text = 'Integer'
        elif type(value) is dict:
            new_cells[1].text = 'Array'
        elif type(value) is bool:
            new_cells[1].text = 'Boolean'
        else:
            new_cells[1].text = 'String'
        if key.find('Action') > -1 or key.find('target') > -1 or key.find('Reset') > -1:
            new_cells[2].text = 'False'
        else:
            new_cells[2].text = 'True'

        #print(key,"   ---------------    ",value)
        if key != "BiosVersion" and value is not None and isinstance(value, str):
            new_cells[3].text = str(value)
    return True 

def flatten (mol, next):
    if isinstance(mol,list):
        for i in mol:
            if isinstance(i,list):
                flatten(i,next)
            elif isinstance(i, dict):
                #print('list value is ', i)
                flatten(i,next)

    elif isinstance(mol, dict):
        for k, v in mol.items():
            if k.find('@odata.id') > -1 or (type(v) is str and v.find('/redfish/v1') > -1):
                #print('#########   ',v , '   ########')
                next.add(v)
            elif isinstance(v, dict):
                flatten(v,next)


def is_json(myjson):
  try:
    json_object = json.loads(myjson)
  except ValueError as e:
    return False
  return True 

def go (root, url):
    if url in visited:
        return
    elif url.find ('redfish') == -1:
        return
    #print(url)
    visited.add(url)
    uri=root+url
    try:
        res = requests.get(str(uri), auth=('Administrator','superuser'), verify=False, timeout=120)
        if res.text is None or not is_json(res.text):
            return
        element=json.loads(res.text)
        next = set()
        for (k, v) in element.items():
            #print(k , ' : ', v)
            if k.find('@odata.id') > -1:
                next.add(v)
            elif isinstance(v,list) or isinstance(v,dict):
                flatten(v, next)

        #print('=============   ',uri, '    =============')
        #print(json.dumps(element, indent=4, sort_keys=True))
        if url not in st.keys():
            print ('***   add in ', url)
            #print element
            st[url] = element
            sts[url] = url.split('/')
        if write2Document(uri, element,next):
            print('$$$$   ',url)
            
        for urll in next:
            #print(urll)
            go (root, urll)
    except requests.exceptions.RequestException as e:
                # catastrophic error. bail.
            print(e)

    

def readDoc(spec):
    #spec = Document('/home/ben/Desktop/amidoc.docx')
    #spec = docx.Document('/home/ben/Desktop/BMC_Nokia_Redfish_API_v0.4.docx')
    """for block in iter_block_items(spec):
        #print(type(block))
        if type(block) is docx.text.paragraph.Paragraph or type(block) is docx.table.Table:
            print(block.text)"""
    #iter_block_items(spec)
    """for para in spec.paragraphs:
        #if para.text.find("https://{{ip}}/") > -1:
            print(para.text)
    for section in spec.sections:
        for table in section.footer.tables:
            for i in range(len(table.rows)):
                for j in range(len(table.columns)):
                    print(table.cell(i,j).text)"""
    #print(len(spec.tables))
    found = True 
    for table in spec.tables:
        #print(len(table.rows))
        #print(len(table.columns))
        found = False 
        for i in range(len(table.rows)):
            for j in range(len(table.columns)):
                #print('####   ',table.cell(i,j).text)
                if table.cell(i,j).text.find ('Type URI') > -1:
                    print('####   ',table.cell(i,j+1).text)
                    dts[table.cell(i,j+1).text] = table.cell(i,j+1).text.split('/')
                    dt[table.cell(i,j+1).text] = table 
                    found = True 
                    break

            if found:
                break

def cleanTypeURI(spec):
    for table in spec.tables:
        for i in range(len(table.rows)):
            if table.cell(i,0).text.find ('Type URI') > -1:
                row = table.rows[i]
                remove_row(table, row)
                break


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)



def iter_block_items(parent):
    if isinstance(parent, docx.document.Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            print('### text is : ',child)
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            print('--------------- table man -----------------')
            for row in table.rows:
                for cell in row.cells:
                    yield iter_block_items(cell)

def find (target):
    length = len(target)
    for uri, urii in dts.items():
        if len(urii) == length:
            found = True 
            for i in range(length):
                if urii[i] != 'benchin' and target[i]!=urii[i]:
                    found = False 
                    break 
            if found:
                return str(uri )

    return "not here"

def main():
    print("pls enter your server ip ...")
    ip=sys.stdin.readline().strip('\n')
    print("pls enter your spec file path and file name ...")
    fp=sys.stdin.readline().strip('\n')
#    ip='10.10.12.11'
#    fp='/home/ben/Desktop/redfish-spec-generator/BMC_Redfish_RTP1.7-API_v0.1.0.docx'
    spec = docx.Document(fp)
    readDoc(spec)
    filenamelist = fp.split('/')
    filename = filenamelist[-1]
    print(filename)
    length = len (filenamelist)
    length = length-1
    print('length is ', length)
    root="https://" + str(ip)
    url="/redfish/v1/"
    filepath=""
    for i in range(length):
        filepath+=filenamelist[i]
        filepath+="/"
    print('filepath is ', filepath)
    print("we start from ",str(root))
    #res=requests.get(root, auth=('Administrator','superuser'))
    #res = requests.get('https://10.10.12.115/redfish/v1/', auth=('admin', 'cmb9.admin'))
    go (str(root),str(url))
    print("document table len:  ",len(dt)," server table length : ",len(st))
    for key in st:
        enter = False 
        dkey="not here"
        if key in dt:
            enter = True 
            dkey = key 
        else: 
            dkey = find (sts[key])
            if dkey != "not here":
                enter = True 
            #iterate this table then save it property as map
        if enter:
            print ( key,"   :    " , dkey )
            property=set()
            print(type(dt[dkey]))
            for row in dt[dkey].rows:
                if len(row.cells):
                    word = str(row.cells[0].text)
                    pos = word.find('(') 
                    if pos != -1:
                        word = word[:pos]
                    """pos = word.find('%') 
                    if pos != -1:
                        word = word[:pos]"""
                    word = word.replace(' ','')
                    property.add(word.lower())
                    print('====     ',word)
            # below is for when spec has no such property then we will append (x) behind that property
            """print(type(st[key]))
            for p in property:
                if p not in st[key].keys():
                    for row in dt[dkey].rows:
                        if len(row.cells) and row.cells[0].text == p: 
                            run = row.cells[0].paragraphs[0].add_run('(X)')
                            #print row.cells[0].text 
                            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                            break"""

            for properties in st[key].keys():
                print ('#####    ', properties)
                p = properties.lower().replace(' ','')
                position = p.find('@') 
                if position > 0:
                    p = p[:position]
                if p not in property:
                    #print(dt[dkey].cell(-1,-1).paragraphs[0])
                    para = dt[dkey].cell(-1,-1).paragraphs[0]
                    r = dt[dkey].add_row()
                    #r.cells[0].text = properties
                    #r.cells[1].text = 'benchin'
                    r1 = r.cells[0].paragraphs[0].add_run(properties)
                    if type(properties) is list:
                        r0 = r.cells[1].paragraphs[0].add_run('Object')
                        r0.font.size = r1.font.size = Pt(8)
                    elif type(properties) is int:
                        r0 = r.cells[1].paragraphs[0].add_run('Integer')
                        r0.font.size = r1.font.size = Pt(8)
                    elif type(properties) is dict:
                        r0 = r.cells[1].paragraphs[0].add_run('Object')
                        r0.font.size = r1.font.size = Pt(8)
                    elif type(properties) is bool:
                        r0 = r.cells[1].paragraphs[0].add_run('Boolean')
                        r0.font.size = r1.font.size = Pt(8)
                    else:
                        r0 = r.cells[1].paragraphs[0].add_run('String')
                        r0.font.size = r1.font.size = Pt(8)
                    r1 = r.cells[2].paragraphs[0].add_run('quanta')

                    if len(para.runs):
                        r0.style = para.runs[0].style
                    print('@@@@@   ', properties, '  not exists')
                """else:
                    for row in dt[dkey].rows:
                        if len(row.cells) and row.cells[0].text == properties: 
                            run = row.cells[0].paragraphs[0].add_run('%%')
                            #print row.cells[0].text 
                            #property.add(row.cells[0].text)
                            run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
                            break"""

        else:
            mdoc.add_paragraph(key, style = 'ListBullet')
            print('cant find  ', key)
    """for key, value in dt.items():
        print(key, " : ", value)
    print('=======================================================')
    for key, value in st.items():
        print(key, " : ", value)"""
    #res=requests.get('https://10.10.12.115/redfish/v1/', auth=('Administrator','superuser'), verify=False)
   # print(res.text)
    missuri = filepath+"missingURI.docx"
    newfile = filepath+"modifiedSpec.docx"
    generateS = filepath+"generateSpec.docx"
    print(missuri, newfile, generateS)
    cleanTypeURI(spec)
    document.save(generateS)
    mdoc.save(missuri)
    spec.save(newfile)
if __name__ == '__main__':
    main()
